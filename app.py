import streamlit as st
import io
import requests
import os
import traceback
from urllib.parse import urlparse

# --- 导入处理 ---
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml  
    from docx.oxml.ns import qn, nsmap 
    
    import mistletoe
    from mistletoe import block_token, span_token
    from mistletoe.base_renderer import BaseRenderer
    
    # 注册命名空间 (防止 KeyError: 'v')
    nsmap['v'] = 'urn:schemas-microsoft-com:vml'
    nsmap['o'] = 'urn:schemas-microsoft-com:office:office'
    
except ImportError as e:
    st.error(f"🚨 依赖库导入失败: {e}")
    st.stop()

# ==========================================
# 工具函数
# ==========================================
def get_available_font():
    """获取系统可用的中文字体（兼容跨平台）"""
    try:
        import win32api
        fonts = [font for font in win32api.GetFontFamilyNames() if '微软雅黑' in font]
        return fonts[0] if fonts else '宋体'
    except:
        # 非Windows系统 fallback
        return 'SimHei' if os.name == 'posix' else '宋体'

def setup_page_layout(doc, image_stream=None):
    """
    统一设置页面布局：A4尺寸、精确边距、背景图
    """
    # 1. 准备背景图关联 (如果有)
    bg_rId = None
    if image_stream:
        try:
            image_stream.seek(0)
            # 图片大小校验
            if image_stream.getbuffer().nbytes > 10 * 1024 * 1024:  # 10MB
                st.warning("⚠️ 背景图超过10MB，文档体积可能过大")
            bg_rId, _ = doc.part.get_or_add_image(image_stream)
        except Exception as e:
            st.error(f"背景图处理失败: {e}")

    # 2. 遍历所有章节进行设置
    for section in doc.sections:
        # --- A. 设置 A4 尺寸 ---
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        
        # --- B. 设置精确边距 ---
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

        # --- C. 设置背景图 (VML) ---
        if bg_rId:
            section_element = section._sectPr
            
            # 构造 VML XML（优化格式）
            vmldata = (
                '<v:background id="_x0000_s1025" o:bwmode="white" fillcolor="white [3212]" '
                'xmlns:v="urn:schemas-microsoft-com:vml" '
                'xmlns:o="urn:schemas-microsoft-com:office:office">'
                f'<v:fill r:id="{bg_rId}" type="frame"/>'
                '</v:background>'
            )
            
            # 清除旧背景
            existing_bg = section_element.find(qn('v:background'))
            if existing_bg is not None:
                section_element.remove(existing_bg)
            
            # 插入新背景
            try:
                bg_element = parse_xml(vmldata)
                section_element.insert(0, bg_element)
            except Exception as e:
                st.warning(f"背景图XML解析失败: {e}")

# ==========================================
# Markdown 渲染器（兼容所有 mistletoe 版本）
# ==========================================
class DocxRenderer(BaseRenderer):
    def __init__(self, doc):
        self.doc = doc
        self.font_name = get_available_font()  # 动态获取字体
        
        # 设置默认样式（正文）
        normal_style = doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = self.font_name
        normal_font.size = Pt(10.5)
        normal_font._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        normal_style.paragraph_format.line_spacing = 1.5  # 正文行间距1.5倍
        
        # 自定义标题样式（区分不同级别）
        self.setup_heading_styles()
        
        super().__init__()

    def setup_heading_styles(self):
        """自定义标题样式，区分标题和正文"""
        # 一级标题
        h1_style = self.doc.styles.add_style('Custom Heading 1', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = self.font_name
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色
        h1_font._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 左对齐
        
        # 二级标题
        h2_style = self.doc.styles.add_style('Custom Heading 2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = self.font_name
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 76, 153)  # 中蓝色
        h2_font._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(9)
        h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 三级标题
        h3_style = self.doc.styles.add_style('Custom Heading 3', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = self.font_name
        h3_font.size = Pt(12)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 102, 204)  # 浅蓝色
        h3_font._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        h3_style.paragraph_format.space_before = Pt(15)
        h3_style.paragraph_format.space_after = Pt(6)
        h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def render_document(self, token):
        if hasattr(token, 'children') and token.children:
            for child in token.children:
                self.render(child)

    def render_heading(self, token):
        """渲染标题，应用自定义样式"""
        level = token.level
        text = self.render_inner(token)
        p = self.doc.add_paragraph(text)
        
        # 根据级别应用不同样式
        if level == 1:
            p.style = self.doc.styles['Custom Heading 1']
        elif level == 2:
            p.style = self.doc.styles['Custom Heading 2']
        elif level >= 3:
            p.style = self.doc.styles['Custom Heading 3']
        
        # 标题文字单独设置（防止样式失效）
        for run in p.runs:
            run.font.name = self.font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)

    def render_paragraph(self, token):
        """渲染正文段落，区分标题样式"""
        paragraph = self.doc.add_paragraph()
        paragraph.style = self.doc.styles['Normal']  # 强制应用正文样式
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        self.render_inner(token, paragraph)

    def render_raw_text(self, token, parent_paragraph=None):
        content = token.content
        if parent_paragraph:
            run = parent_paragraph.add_run(content)
            run.font.name = self.font_name
            run.font.size = Pt(10.5)  # 正文固定字号
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            return run
        return content

    def render_strong(self, token, parent_paragraph):
        run = self.render_inner(token, parent_paragraph)
        if run: 
            run.bold = True

    def render_emphasis(self, token, parent_paragraph):
        run = self.render_inner(token, parent_paragraph)
        if run: 
            run.italic = True
        
    def render_list(self, token):
        """
        兼容所有 mistletoe 版本的列表渲染
        - 旧版本：List 类 + start 属性
        - 新版本：OrderedList/UnorderedList 类
        """
        # 判定是否为有序列表（终极兼容方案）
        is_ordered = False
        
        # 适配新版本 mistletoe（有 OrderedList 类）
        if hasattr(block_token, 'OrderedList'):
            is_ordered = isinstance(token, block_token.OrderedList)
        # 适配旧版本 mistletoe（只有 List 类，通过 start 属性判断）
        else:
            is_ordered = hasattr(token, 'start') and token.start is not None
        
        # 设置列表样式（有序=数字，无序=圆点）
        list_style = 'List Number' if is_ordered else 'List Bullet'
        
        # 渲染列表项
        if hasattr(token, 'children') and token.children:
            for list_item in token.children:
                if hasattr(list_item, 'children') and list_item.children:
                    paragraph = self.doc.add_paragraph(style=list_style)
                    # 列表项应用正文样式
                    paragraph.style = self.doc.styles['Normal']
                    # 递归渲染列表项内容（支持嵌套）
                    self.render_list_item(list_item, paragraph)

    def render_list_item(self, token, parent_paragraph=None):
        """处理列表项（兼容所有版本，支持嵌套）"""
        if not parent_paragraph:
            parent_paragraph = self.doc.add_paragraph()
            parent_paragraph.style = self.doc.styles['Normal']
        
        for child in token.children:
            # 兼容所有版本的列表类判断
            list_classes = [block_token.List]
            if hasattr(block_token, 'OrderedList'):
                list_classes.append(block_token.OrderedList)
            if hasattr(block_token, 'UnorderedList'):
                list_classes.append(block_token.UnorderedList)
            
            if isinstance(child, tuple(list_classes)):
                self.render_list(child)  # 递归渲染嵌套列表
            else:
                self.render_inner(child, parent_paragraph)

    def render_image(self, token, parent_paragraph):
        url = token.src.strip()
        alt_text = token.title if token.title else "图片"
        image_stream = None

        try:
            if not url:
                raise ValueError("图片URL为空")
            
            # 区分网络/本地图片
            parsed = urlparse(url)
            if parsed.scheme in ('http', 'https'):
                response = requests.get(url, timeout=10, stream=True)
                response.raise_for_status()
                image_stream = io.BytesIO(response.content)
                response.close()
            elif os.path.exists(url):
                with open(url, 'rb') as f:
                    image_stream = io.BytesIO(f.read())
            else:
                raise FileNotFoundError(f"图片路径无效: {url}")

            # 插入图片（限制宽度）
            run = parent_paragraph.add_run()
            max_width = Cm(min(14, self.doc.sections[0].page_width.cm - 4))
            run.add_picture(image_stream, width=max_width)
            # 图片说明
            caption = parent_paragraph.add_run(f"\n{alt_text}")
            caption.italic = True
            caption.font.name = self.font_name
            caption.font.size = Pt(9)

        except Exception as e:
            err_run = parent_paragraph.add_run(f"[图片加载失败: {alt_text} - {str(e)}]")
            err_run.font.color.rgb = RGBColor(255, 0, 0)
        finally:
            if image_stream:
                image_stream.close()

    def render_table(self, token):
        """渲染表格，设置居中+美化样式"""
        if not hasattr(token, 'children') or not token.children: 
            return
        rows = len(token.children)
        if rows == 0: 
            return
        
        first_row = token.children[0]
        if not hasattr(first_row, 'children') or not first_row.children: 
            return
        cols = len(first_row.children)
        
        # 创建表格并设置居中
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
        table.autofit = True  # 自动适配宽度
        
        # 设置表头样式
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 表头文字居中
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # 白色文字
                run.font.name = self.font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            # 表头背景色（深蓝色）
            cell_shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="003366"/>')
            cell._tc.get_or_add_tcPr().append(cell_shading)
        
        # 填充表格内容并设置居中
        for i, row_token in enumerate(token.children):
            row = table.rows[i]
            # 设置行高
            row.height = Pt(24)
            row.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.AT_LEAST
            
            if hasattr(row_token, 'children') and row_token.children:
                for j, cell_token in enumerate(row_token.children):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        cell._element.clear_content()
                        paragraph = cell.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 单元格文字居中
                        self.render_inner(cell_token, paragraph)
                        
                        # 设置单元格字体
                        for run in paragraph.runs:
                            run.font.name = self.font_name
                            run.font.size = Pt(10)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
                            
                            # 非表头行设置浅灰色背景
                            if i > 0:
                                if j % 2 == 0:
                                    cell_shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F5F5F5"/>')
                                    cell._tc.get_or_add_tcPr().append(cell_shading)

    def render_inner(self, token, parent_paragraph=None):
        if hasattr(token, 'children') and token.children:
            last_run = None
            for child in token.children:
                if isinstance(child, span_token.RawText):
                    last_run = self.render_raw_text(child, parent_paragraph)
                elif isinstance(child, span_token.Strong):
                    self.render_strong(child, parent_paragraph)
                elif isinstance(child, span_token.Emphasis):
                    self.render_emphasis(child, parent_paragraph)
                elif isinstance(child, span_token.Image):
                    self.render_image(child, parent_paragraph)
            return last_run
        elif hasattr(token, 'content'):
            return self.render_raw_text(token, parent_paragraph)
        return None

# ==========================================
# 界面逻辑
# ==========================================
st.set_page_config(page_title="Huamai 文档生成器", layout="wide", page_icon="📄")
st.title("📄 Huamai 文档生成工具 (V5.3 样式优化版)")

col1, col2 = st.columns([4, 6])

with col1:
    st.info("💡 请上传 A4 背景图（建议大小≤10MB）")
    bg_file = st.file_uploader("上传背景图 (PNG/JPG)", type=['png', 'jpg', 'jpeg'])
    generate_btn = st.button("🚀 生成文档", type="primary", use_container_width=True)

with col2:
    default_md = """# 产品规格说明书

## 1. 简介
本产品完全符合 A4 打印标准，页边距已严格校准，标题和正文样式区分明显，表格支持居中显示。

## 2. 详细参数
| 项目 | 规格 | 说明 |
| :--- | :--- | :--- |
| 尺寸 | A4 | 标准纸张 |
| 边距 | 定制 | 72/72/54/54 pt |
| 标题1 | 16号字 | 深蓝色加粗 |
| 标题2 | 14号字 | 中蓝色加粗 |
| 正文 | 10.5号字 | 常规样式 |

## 3. 功能列表
- 支持Markdown语法
  - 多级标题（样式区分）
  - 粗体、斜体
  - 有序/无序列表
- 表格居中显示
  - 表头背景色
  - 单元格文字居中
  - 奇偶行隔行变色
- 图片插入（网络/本地）
- 自定义背景图

## 4. 有序列表示例
1. 第一步：输入Markdown内容
2. 第二步：上传背景图（可选）
3. 第三步：点击生成文档

## 5. 示例图片
![示例图片](https://picsum.photos/800/600)
"""
    md_input = st.text_area("Markdown 内容", height=500, value=default_md)

if generate_btn:
    if not md_input.strip():
        st.error("❌ 请输入 Markdown 内容！")
    else:
        with st.spinner("正在排版文档..."):
            try:
                # 初始化文档
                doc = Document()
                
                # 渲染Markdown内容
                renderer = DocxRenderer(doc)
                doc_token = mistletoe.Document(md_input)
                renderer.render(doc_token)
                
                # 处理背景图
                bg_stream = None
                if bg_file:
                    bg_stream = io.BytesIO(bg_file.getvalue())
                setup_page_layout(doc, bg_stream)
                
                # 保存文档到IO流
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                # 关闭临时流
                if bg_stream:
                    bg_stream.close()
                
                st.success("✅ 文档生成成功！（已优化标题样式和表格居中）")
                st.download_button(
                    label="📥 下载最终文档",
                    data=doc_io,
                    file_name="Huamai_Styled_Final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"❌ 生成失败: {e}")
                st.expander("查看详细错误信息").code(traceback.format_exc())
